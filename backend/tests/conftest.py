from __future__ import annotations

import io
from pathlib import Path
from typing import Any

import fitz
import pytest
from docx import Document
from PIL import Image, ImageDraw, ImageFont

from app.config import Settings
from app.services.ocr.image_ocr import OcrTextResult


@pytest.fixture()
def settings(tmp_path: Path) -> Settings:
    return Settings(
        _env_file=None,
        ocr_max_upload_mb=5,
        ocr_max_pages=3,
        ocr_languages="vie+eng",
        ocr_temp_ttl_seconds=0,
        ocr_report_dir=tmp_path / "artifacts" / "reports",
        ocr_log_file=tmp_path / "artifacts" / "logs" / "events.jsonl",
        ocr_runtime_dir=tmp_path / "runtime",
        gemini_api_key="",
        cors_allowed_origins=["http://localhost:8000"],
    )


def _font(size: int = 28) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    path = Path("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf")
    return ImageFont.truetype(str(path), size) if path.exists() else ImageFont.load_default()


@pytest.fixture()
def image_bytes() -> bytes:
    image = Image.new("RGB", (1400, 900), "white")
    draw = ImageDraw.Draw(image)
    draw.text((60, 60), "Ky nang: Python, SQL, FastAPI", fill="black", font=_font())
    draw.text((60, 120), "Du an: Cong cu phan tich du lieu bang Python", fill="black", font=_font())
    output = io.BytesIO()
    image.save(output, format="PNG")
    return output.getvalue()


@pytest.fixture()
def scan_pdf_bytes(image_bytes: bytes) -> bytes:
    document = fitz.open()
    page = document.new_page(width=595, height=842)
    page.insert_image(page.rect, stream=image_bytes, keep_proportion=True)
    data = document.tobytes()
    document.close()
    return data


@pytest.fixture()
def text_pdf_bytes() -> bytes:
    document = fitz.open()
    page = document.new_page()
    text = (
        "Ky nang: Python, SQL, FastAPI. "
        "Du an: Xay dung ung dung phan tich du lieu bang Python va FastAPI. "
        "Kinh nghiem duoc nguoi dung xac nhan truoc khi de xuat de tai."
    )
    page.insert_text((72, 100), text, fontsize=11)
    data = document.tobytes()
    document.close()
    return data


@pytest.fixture()
def docx_bytes() -> bytes:
    document = Document()
    document.add_heading("Hồ sơ dự án tổng hợp", level=1)
    document.add_paragraph("Kỹ năng: Python, SQL, FastAPI")
    document.add_paragraph("Dự án: Dashboard học tập")
    document.add_paragraph("Xây dựng API bằng FastAPI và phân tích dữ liệu bằng Python, SQL.")
    document.add_paragraph("Sở thích: Dữ liệu, AI và giáo dục")
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


class FakeImageOcr:
    def __init__(
        self,
        text: str = "Kỹ năng: Python, SQL\nDự án: OCR local\nXây dựng bằng Python và FastAPI.",
        confidence: float = 88.0,
        warnings: list[str] | None = None,
    ) -> None:
        self.text = text
        self.confidence = confidence
        self.warnings = warnings or []
        self.calls = 0

    def extract(self, _image: Image.Image, *, languages: str, on_event: Any = None) -> OcrTextResult:
        self.calls += 1
        if on_event:
            on_event("ocr_started", "started", {"extraction_method": "tesseract_ocr"})
            on_event(
                "ocr_completed",
                "success",
                {
                    "extraction_method": "tesseract_ocr",
                    "character_count": len(self.text),
                    "ocr_confidence": self.confidence,
                    "retry_performed": False,
                },
            )
        return OcrTextResult(
            text=self.text,
            confidence=self.confidence,
            language=languages,
            warnings=self.warnings,
        )


@pytest.fixture()
def fake_ocr() -> FakeImageOcr:
    return FakeImageOcr()
