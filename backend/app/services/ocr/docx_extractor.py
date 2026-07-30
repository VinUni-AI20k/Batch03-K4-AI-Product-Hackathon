from __future__ import annotations

from pathlib import Path

from docx import Document

from .errors import OcrPipelineError
from .models import ExtractionResult, PageText


class DocxExtractor:
    def extract(self, path: Path) -> ExtractionResult:
        try:
            document = Document(path)
            blocks: list[str] = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
            for table in document.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if cells:
                        blocks.append(" | ".join(cells))
        except (OSError, ValueError, KeyError) as exc:
            raise OcrPipelineError("CORRUPTED_FILE", "DOCX text extraction failed.") from exc

        text = "\n".join(blocks)
        return ExtractionResult(
            pages=[PageText(page=1, text=text, source_type="docx_text")],
            primary_method="python_docx",
            ocr_used=False,
            text_pages=[1],
        )
